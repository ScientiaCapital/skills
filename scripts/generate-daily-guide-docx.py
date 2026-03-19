#!/usr/bin/env python3
"""Generate Tim's Daily Operations Guide as DOCX for Google Drive upload."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import os

OUTPUT = os.path.expanduser("~/Desktop/Tim_Daily_Operations_Guide.docx")
TODAY = datetime.now().strftime("%B %d, %Y")

# Colors
PRIMARY = RGBColor(0x0f, 0x34, 0x60)
ACCENT = RGBColor(0xe9, 0x45, 0x60)
DARK = RGBColor(0x1a, 0x1a, 0x2e)
GRAY = RGBColor(0x6b, 0x72, 0x80)
WHITE = RGBColor(0xff, 0xff, 0xff)
LIGHT_BG = RGBColor(0xf0, 0xf4, 0xf8)

def set_cell_bg(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
        set_cell_bg(cell, '0f3460')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = DARK
            if r_idx % 2 == 1:
                set_cell_bg(cell, 'f0f4f8')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # ─── PAGE 1: COVER ───
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Tim's Daily Operations Guide")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("BDR Playbook  |  Epiphan Video")
    run.font.size = Pt(14)
    run.font.color.rgb = GRAY

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"Generated: {TODAY}  |  67 Skills  |  7 MCP Servers  |  90+ Tools")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    doc.add_paragraph()

    h = doc.add_heading('Quick Start', level=2)
    for run in h.runs:
        run.font.color.rgb = PRIMARY

    p = doc.add_paragraph()
    p.add_run('Say ').font.size = Pt(10)
    r = p.add_run('"morning brief"')
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(10)
    p.add_run(' to start your day.').font.size = Pt(10)

    add_styled_table(doc,
        ["Command", "What It Does"],
        [
            ['"morning brief"', "Run all morning workflows"],
            ['"deal momentum"', "Score pipeline, flag RED/YELLOW deals"],
            ['"portfolio update"', "Attribute closed deals to skills"],
            ['"market digest"', "Pre-market scan + IBKR positions"],
            ['"prep me for [Company]"', "Auto-generate MEDDIC call prep"],
            ['"prospect [Company]"', "End-to-end research + cadence"],
        ],
        [2, 4.5]
    )

    p = doc.add_paragraph()
    r = p.add_run("March 2026 Targets (Ramp 50%): 12 deals | $357K pipeline | $125K revenue | 50+ daily dials")
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(10)

    # ─── PAGE 2: TIMELINE ───
    doc.add_page_break()
    h = doc.add_heading('Morning Workflow Timeline', level=1)
    for run in h.runs:
        run.font.color.rgb = PRIMARY

    h2 = doc.add_heading('MONDAY (Full Automation Chain)', level=2)
    add_styled_table(doc,
        ["Time", "Skill", "What Happens"],
        [
            ["06:00", "prospect-enrich", "Phoneless -> Apollo/Clay enrichment (DEMO REQUEST first)"],
            ["06:15", "phone-waterfall", "Verify phones -> callable queue"],
            ["06:30", "prospect-refresh", "Net-new ICP search (7 verticals) + Gmail drafts"],
            ["07:00", "deal-momentum", "Score deals, flag RED/YELLOW, draft recovery emails"],
            ["07:00", "portfolio-linker", "Link closed deals -> GTME evidence"],
            ["07:00", "trading-alerts", "Pre-market scan + IBKR positions"],
            ["07:15", "sequence-load", "Auto-enroll prospects -> Apollo sequences"],
            ["07:25", "callable-lead-count", "ATL/BTL inventory health check"],
            ["07:30", "morning-brief", "Calendar + dials + deals + drafts briefing"],
        ],
        [0.6, 1.5, 4.4]
    )

    h2 = doc.add_heading('DAILY (Tuesday - Friday)', level=2)
    add_styled_table(doc,
        ["Time", "Skill", "What Happens"],
        [
            ["06:00", "prospect-enrich", "Daily phoneless enrichment"],
            ["07:00", "deal-momentum + portfolio-linker + trading-alerts", "Pipeline + portfolio + market"],
            ["07:25", "callable-lead-count", "ATL/BTL inventory + runway"],
            ["07:30", "morning-brief", "Full daily briefing"],
        ],
        [0.6, 3.2, 2.7]
    )

    # ─── PAGE 3: TRIGGERS ───
    doc.add_page_break()
    h = doc.add_heading('Trigger Phrase Quick Reference', level=1)
    for run in h.runs:
        run.font.color.rgb = PRIMARY

    categories = {
        "Core (5)": [
            ["workflow-orchestrator", '"start day", "begin session"'],
            ["cost-metering", '"cost check", "budget status"'],
            ["portfolio-artifact", '"portfolio report", "what did I ship"'],
            ["project-context", '"load project context", "end session"'],
            ["workflow-enforcer", "Automatic on all sessions"],
        ],
        "BDR Automation (5)": [
            ["prospect-enrich", '"run prospect enrich", "enrich phoneless"'],
            ["prospect-refresh", '"run prospect refresh", "search new ICP"'],
            ["sequence-load", '"load sequences", "enroll new leads"'],
            ["callable-lead-count", '"show callable leads", "ATL runway"'],
            ["morning-brief", '"morning brief", "today\'s dial list"'],
        ],
        "Sales Workflow (6)": [
            ["prospect-research-to-cadence", '"research prospect", "prospect [Company]"'],
            ["phone-verification-waterfall", '"verify phones", "dial list"'],
            ["meddic-call-prep-auto", '"call prep", "prep me for [Company]"'],
            ["deal-momentum-analyzer", '"deal health", "pipeline review"'],
            ["portfolio-deal-linker", '"portfolio update", "gtme evidence"'],
            ["trading-alert-scheduler", '"market digest", "pre-market scan"'],
        ],
        "Sales Intelligence (13)": [
            ["champion-identifier", '"find champion", "internal champion"'],
            ["cold-email-sequence-gen", '"cold email sequence", "email campaign"'],
            ["contact-hunter", '"find contact info", "who works at"'],
            ["email-template-generator", '"email template", "write email"'],
            ["inbound-lead-qualifier", '"qualify inbound", "score lead"'],
            ["intent-signal-aggregator", '"intent signals", "buyer intent"'],
            ["linkedin-sales-nav-alt", '"linkedin prospects", "sales navigator"'],
            ["lookalike-customer-finder", '"lookalike companies", "target accounts"'],
            ["meeting-intelligence", '"meeting notes", "transcript analysis"'],
            ["personalization-at-scale", '"personalize at scale", "first lines"'],
            ["pipeline-health-analyzer", '"pipeline health", "forecast accuracy"'],
            ["sales-methodology-impl", '"sales methodology", "implement MEDDIC"'],
            ["social-selling-content", '"social selling", "LinkedIn content"'],
        ],
        "Strategy (5)": [
            ["business-model-canvas", '"business model", "canvas"'],
            ["blue-ocean-strategy", '"blue ocean", "ERRC"'],
            ["jobs-to-be-done", '"JTBD", "customer jobs"'],
            ["challenger-sale", '"challenger sale", "teach tailor take control"'],
            ["never-split-the-difference", '"negotiation", "tactical empathy"'],
        ],
    }

    for cat, skills in categories.items():
        doc.add_heading(cat, level=2)
        add_styled_table(doc, ["Skill", "Trigger Phrases"], skills, [2.2, 4.3])

    # ─── PAGE 4-5: COMBOS ───
    doc.add_page_break()
    h = doc.add_heading('Skill Combinations by Use Case', level=1)
    for run in h.runs:
        run.font.color.rgb = PRIMARY

    add_styled_table(doc,
        ["Use Case", "Trigger", "Chain", "Impact"],
        [
            ["Start of Day", '"morning brief"', "prospect-enrich -> phone-waterfall -> prospect-refresh -> sequence-load -> callable-lead-count -> morning-brief", "Full pipeline automation"],
            ["Prospect a Company", '"prospect [Co]"', "prospect-research-to-cadence -> Apollo/Clay -> Gmail drafts -> Apollo sequence", "90 min/day saved"],
            ["Call Prep", '"prep me for [Co]"', "meddic-call-prep-auto -> HubSpot + Apollo + Clari", "60-sec MEDDIC brief"],
            ["Pipeline Review", '"deal momentum"', "deal-momentum-analyzer -> 6-signal scoring -> Gmail recovery drafts", "Recover stalled deals"],
            ["Research + Outreach", '"research prospect"', "prospect-research-to-cadence -> cold-email-sequence-gen -> personalization-at-scale", "Full pipeline + personalized emails"],
            ["Trading", '"market digest"', "trading-alert-scheduler -> trading-signals + ibkr-api", "3-min pre-market read"],
            ["End of Day", '"portfolio report"', "portfolio-artifact -> portfolio-deal-linker -> cost-metering", "Capture metrics + costs"],
            ["Strategy", '"JTBD analysis"', "JTBD -> blue-ocean -> challenger-sale -> NSTTD", "Full strategy chain"],
            ["Debug", '"debug this"', "debug-like-expert -> hypothesis testing", "Systematic investigation"],
            ["Feature Dev", '"plan this"', "planning-prompts -> agent-teams -> testing -> git-workflow", "Plan -> parallel dev -> TDD"],
        ],
        [1.1, 1.2, 2.6, 1.6]
    )

    # ─── PAGE 6: MCP ───
    doc.add_page_break()
    h = doc.add_heading('MCP Tools Available', level=1)
    for run in h.runs:
        run.font.color.rgb = PRIMARY

    doc.add_paragraph("7 MCP servers providing 90+ live tools for real-time data.")

    add_styled_table(doc,
        ["Server", "Tools", "Key Capabilities", "Used For"],
        [
            ["HubSpot", "14", "Deals, contacts, companies, CRM search", "Sales workflow, pipeline"],
            ["Apollo.io", "15", "Enrich, search, contacts, sequences", "Prospecting, outreach"],
            ["Clay", "13+", "Waterfall enrichment, contact search", "Phone waterfall"],
            ["Gmail", "7", "Drafts, search, read, labels", "Email drafts"],
            ["Google Calendar", "9", "Events, free time, meetings", "Call prep, scheduling"],
            ["IBKR", "7", "Portfolio, account, margin, shorts", "Trading, positions"],
            ["Supabase", "20+", "SQL, migrations, edge functions", "Disposition tracking"],
        ],
        [1, 0.5, 2.8, 2.2]
    )

    doc.add_heading('Golden Rules (Auto-Enforced via Hooks)', level=2)
    add_styled_table(doc,
        ["Rule", "Check", "Action"],
        [
            ["Customer Check", "CRM lookup before enrichment", "DENY if existing customer"],
            ["Device Owner", "device_count >= 1", "DENY enrichment"],
            ["Channel Partner", "is_channel = true", "DENY enrichment"],
            ["AE-Owned", "hubspot_owner_id IN (82625923, 423155215)", "DENY prospecting"],
            ["Product Setup", 'first_conversion contains "setup"', "DENY"],
            ["DEMO REQUEST", 'first_conversion contains "demo", "pricing"', "PRIORITIZE (highest tier)"],
        ],
        [1.3, 2.5, 2.7]
    )

    # ─── PAGE 7: TARGETS ───
    doc.add_page_break()
    h = doc.add_heading('Daily Targets + ATL/BTL Quick Reference', level=1)
    for run in h.runs:
        run.font.color.rgb = PRIMARY

    doc.add_heading('March 2026 Ramp Targets (50%)', level=2)
    add_styled_table(doc,
        ["Metric", "Monthly Target", "Stretch (126%+)", "Daily Run Rate"],
        [
            ["Deals created", "12", "16+ (1.5x)", "0.6/day"],
            ["Pipeline", "$357,000", "$450,000+", "$17,850/day"],
            ["Revenue", "$125,000", "$157,000+", "$6,250/day"],
            ["Daily dials", "50+", "60+", "50/day min"],
            ["Connect rate", "8-12%", "15%+", "4-6/day"],
            ["Speed to lead", "<60 min", "<30 min", "Immediate"],
        ],
        [1.5, 1.5, 1.5, 2]
    )

    doc.add_heading('ATL/BTL Classification', level=2)
    doc.add_paragraph("ATL = can sign a purchase order. BTL = needs someone else's approval.")

    add_styled_table(doc,
        ["Tier", "Who", "Action"],
        [
            ["DEMO REQUEST", "first_conversion: demo, pricing, contact us", "ENRICH FIRST"],
            ["ATL", "Chief/C-suite, VP, President, Provost, Director, Dean, Court Admin", "Always prospect"],
            ["GRAY ($25K)", "Manager (AV/IT), Dept Chair, Dir Ed Tech", "Flag for review"],
            ["BTL", "Technician, Specialist, Coordinator, Support, Admin, Engineer", "Deprioritize"],
            ["NEVER ATL", "Warehouse Mgr, Network Mgr, AV Tech, Lab Coord, Maintenance", "HARD SKIP"],
        ],
        [1.2, 3, 2.3]
    )

    doc.add_heading('ICP Vertical Scores', level=2)
    add_styled_table(doc,
        ["Vertical", "ICP", "Top ATL Titles", "Budget Floor"],
        [
            ["Higher Ed", "90", "CIO, VP IT, Provost, Dean", ">$50K"],
            ["Courts", "85", "Clerk of Court, Court Admin", ">$75K"],
            ["Government", "80", "City Manager, IT Director", ">$100K"],
            ["Corporate AV", "80", "VP Facilities, VP IT/CIO", ">$150K"],
            ["Healthcare", "75", "CFO, CIO, COO, VP Ops", ">$100K"],
            ["Houses of Worship", "70", "Senior Pastor, Exec Pastor", ">$25K"],
            ["K-12", "65", "Superintendent, CTO", ">$10K"],
        ],
        [1.2, 0.5, 3, 1.8]
    )

    # ─── PAGE 8: CHAINS ───
    doc.add_page_break()
    h = doc.add_heading('Workflow Chain Diagrams', level=1)
    for run in h.runs:
        run.font.color.rgb = PRIMARY

    doc.add_heading('1. BDR Daily Pipeline (Monday)', level=2)
    add_styled_table(doc,
        ["Step", "Time", "Skill", "Output"],
        [
            ["1", "06:00", "prospect-enrich", "Enriched contacts (DEMO REQ -> ATL -> GRAY -> BTL)"],
            ["2", "06:15", "phone-waterfall", "Phones verified via Apollo -> Clay"],
            ["3", "06:30", "prospect-refresh", "30 net-new prospects + Gmail drafts"],
            ["4", "07:00", "deal-momentum + portfolio + trading", "Pipeline + deals + market"],
            ["5", "07:15", "sequence-load", "Prospects enrolled in Apollo"],
            ["6", "07:25", "callable-lead-count", "ATL runway + health alerts"],
            ["7", "07:30", "morning-brief", "HTML brief: calendar + 20 leads + drafts"],
        ],
        [0.4, 0.6, 2.5, 3]
    )

    doc.add_heading('2. Automated Deal Flow', level=2)
    add_styled_table(doc,
        ["Stage", "Skill", "MCP Tools", "Output"],
        [
            ["Research", "prospect-research-to-cadence", "Apollo, CRM, Gmail", "Enriched prospects + drafts"],
            ["Verify", "phone-verification-waterfall", "Apollo, Clay", "Verified phone queue"],
            ["Prep", "meddic-call-prep-auto", "HubSpot, Apollo, Cal, Clari", "MEDDIC brief"],
            ["Score", "deal-momentum-analyzer", "HubSpot, Clari, Apollo", "6-signal deal score"],
            ["Attribute", "portfolio-deal-linker", "HubSpot, Apollo, Gmail", "GTME evidence"],
        ],
        [0.7, 2.2, 1.8, 1.8]
    )

    doc.add_heading('3. Strategy Chain', level=2)
    add_styled_table(doc,
        ["Step", "Skill", "Framework"],
        [
            ["1", "jobs-to-be-done", "Christensen JTBD + Ulwick ODI"],
            ["2", "blue-ocean-strategy", "ERRC + Strategy Canvas"],
            ["3", "business-model-canvas", "Osterwalder 9 blocks"],
            ["4", "challenger-sale", "Teach-Tailor-Take Control"],
            ["5", "never-split-the-difference", "Voss FBI negotiation"],
        ],
        [0.5, 2.5, 3.5]
    )

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Generated {TODAY} by Claude Code  |  67 skills")
    r.font.size = Pt(8)
    r.font.color.rgb = GRAY

    doc.save(OUTPUT)
    print(f"DOCX generated: {OUTPUT}")

if __name__ == "__main__":
    main()
